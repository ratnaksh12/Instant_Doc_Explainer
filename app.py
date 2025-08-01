import streamlit as st
import os
import pypdf
import docx
import tempfile
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.docstore.document import Document
from langchain_groq import ChatGroq
from langchain.chains import ConversationalRetrievalChain

# --- APP CONFIGURATION ---
st.set_page_config(page_title="Streamlit Doc Explainer", layout="wide")

# --- GROQ & LLM SETUP ---
# You need to set your Groq API key here.
# It is recommended to use Streamlit's secrets management for this.
# Create a .streamlit/secrets.toml file with: GROQ_API_KEY = "your-api-key"
# For now, you can paste it directly below if needed.
groq_api_key = "your api key"
os.environ['GROQ_API_KEY'] = groq_api_key

EMBED_MODEL = "sentence-transformers/all-MiniLM-L6-v2"
CHAT_MODEL = "meta-llama/llama-4-scout-17b-16e-instruct"

# --- RAG SYSTEM FUNCTIONS ---
def get_file_text(file_data, file_name):
    """Extracts text from various file types from a file-like object."""
    try:
        file_extension = os.path.splitext(file_name)[1].lower()
        
        # Use a temporary file to handle the file-like object
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as temp_file:
            temp_file.write(file_data.read())
            temp_file_path = temp_file.name

        text = ""
        if file_extension == '.pdf':
            reader = pypdf.PdfReader(temp_file_path)
            text = "".join(p.extract_text() for p in reader.pages if p.extract_text())
        elif file_extension == '.docx':
            doc = docx.Document(temp_file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        elif file_extension == '.txt':
            with open(temp_file_path, 'r', encoding='utf-8') as f:
                text = f.read()
        else:
            return ""
        
        os.remove(temp_file_path)
        return text
    except Exception as e:
        st.error(f"Error processing {file_name}: {e}")
        return ""

def create_vector_store(uploaded_files):
    """Creates a FAISS vector store from a list of uploaded files."""
    if not uploaded_files:
        return None

    docs = []
    for uploaded_file in uploaded_files:
        text = get_file_text(uploaded_file, uploaded_file.name)
        if text:
            docs.append(Document(page_content=text, metadata={"source": uploaded_file.name}))
    
    if not docs:
        return None
    
    # Split documents into chunks for better retrieval
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    split_docs = text_splitter.split_documents(docs)
    
    # Create and return the FAISS vector store
    embeddings = HuggingFaceEmbeddings(model_name=EMBED_MODEL)
    vector_store = FAISS.from_documents(split_docs, embeddings)
    return vector_store

def get_qa_chain(vector_store):
    """Initializes the ConversationalRetrievalChain with a Groq LLM."""
    if vector_store is None:
        return None
        
    retriever = vector_store.as_retriever(search_kwargs={"k": 3})
    llm = ChatGroq(model=CHAT_MODEL, api_key=groq_api_key)
    
    # Create the conversational chain
    qa_chain = ConversationalRetrievalChain.from_llm(
        llm,
        retriever,
        return_source_documents=False
    )
    return qa_chain

# --- UI & LOGIC ---
def main():
    st.title("Instant Doc Explainer")
    st.subheader("Upload PDF, DOCX, or TXT files and ask questions about their content.")

    # File uploader section
    uploaded_files = st.sidebar.file_uploader(
        "Upload your documents (PDF, DOCX, TXT)",
        type=["pdf", "docx", "txt"],
        accept_multiple_files=True
    )

    # State management with Streamlit's session_state
    if "qa_chain" not in st.session_state:
        st.session_state.qa_chain = None
    if "messages" not in st.session_state:
        st.session_state.messages = []

    if uploaded_files:
        if "vector_store" not in st.session_state or len(uploaded_files) != len(st.session_state.uploaded_files):
            # Only re-process if files are new or different
            with st.spinner("Processing documents..."):
                st.session_state.vector_store = create_vector_store(uploaded_files)
                st.session_state.uploaded_files = uploaded_files
                if st.session_state.vector_store:
                    st.session_state.qa_chain = get_qa_chain(st.session_state.vector_store)
                    st.success("Documents processed! You can start chatting.")
                else:
                    st.error("Failed to process documents. Please check file formats.")
                    st.session_state.qa_chain = None
    else:
        st.session_state.qa_chain = None
        st.session_state.messages = []
        st.warning("Please upload documents to begin.")

    # Display chat messages from history
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

    # Handle user input
    if question := st.chat_input("Ask a question about your documents..."):
        # Add user message to chat history
        st.session_state.messages.append({"role": "user", "content": question})
        with st.chat_message("user"):
            st.markdown(question)

        # Check if the QA chain is ready
        if st.session_state.qa_chain:
            with st.chat_message("assistant"):
                with st.spinner("Thinking..."):
                    # Format chat history for the chain
                    chat_history = [(m['content'], "") for m in st.session_state.messages if m['role'] == 'user']
                    response = st.session_state.qa_chain.invoke({"question": question, "chat_history": chat_history})
                    answer = response["answer"]
                    st.markdown(answer)

            # Add assistant response to chat history
            st.session_state.messages.append({"role": "assistant", "content": answer})
        else:
            with st.chat_message("assistant"):
                st.warning("Please upload and process documents first.")

if __name__ == "__main__":
    main()
